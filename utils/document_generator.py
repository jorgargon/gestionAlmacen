"""
Document generation utilities for reception forms
"""
import os
import smtplib
import tempfile
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from datetime import datetime
from docx import Document


# Template paths
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_EN = os.path.join(BASE_DIR, 'F1P13_REGISTRO_EN.docx')
TEMPLATE_MP = os.path.join(BASE_DIR, 'F1P13_REGISTRO_MP.docx')

# Output directory for generated documents
OUTPUT_DIR = os.path.join(BASE_DIR, 'generated_docs')
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Email configuration
EMAIL_RECIPIENT = 'jordi@globalispharma.com'
SMTP_SERVER = 'smtp.gmail.com'
SMTP_PORT = 465
EMAIL_USER = 'jordi@zaudera.com'
EMAIL_PASS = 'eaqeslnzskublkwb'


def replace_text_in_paragraph(paragraph, replacements):
    """Replace placeholders in a paragraph while preserving formatting"""
    from docx.shared import Pt
    
    for key, value in replacements.items():
        if key in paragraph.text:
            # Replace in the full paragraph text
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, str(value))
                    # Special formatting for X markers (ENV/MP checkboxes)
                    if (key == '<<ENV>>' or key == '<<MP>>') and value == 'X':
                        run.bold = True
                        if run.font.size:
                            run.font.size = Pt(run.font.size.pt * 2)
                        else:
                            run.font.size = Pt(22)  # Default to 22pt (double of typical 11pt)


def replace_text_in_table(table, replacements):
    """Replace placeholders in all cells of a table"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)


def generate_reception_document(reception_data):
    """
    Generate a filled reception document based on the reception data
    
    Args:
        reception_data: dict with keys:
            - product_type: 'raw_material' or 'packaging'
            - reception_date: date string (YYYY-MM-DD)
            - supplier: string
            - product_name: string
            - quantity: float
            - unit: string
            - units_reviewed: string
            - lot_number: string
            - expiration_date: date string or None
    
    Returns:
        tuple: (docx_path, pdf_path)
    """
    # Select template based on product type
    if reception_data['product_type'] == 'raw_material':
        template_path = TEMPLATE_MP
    else:
        template_path = TEMPLATE_EN
    
    # Load the template
    doc = Document(template_path)
    
    # Format dates
    reception_date = reception_data['reception_date']
    if isinstance(reception_date, str):
        try:
            reception_date = datetime.strptime(reception_date, '%Y-%m-%d').strftime('%d/%m/%Y')
        except:
            pass
    
    expiration_date = reception_data.get('expiration_date')
    if expiration_date:
        if isinstance(expiration_date, str):
            try:
                expiration_date = datetime.strptime(expiration_date, '%Y-%m-%d').strftime('%d/%m/%Y')
            except:
                pass
    else:
        expiration_date = 'SIN CADUCIDAD'
    
    # Build replacements dictionary
    replacements = {
        '<<FECHA_RECEPCION>>': reception_date,
        '<<PROVEEDOR>>': reception_data.get('supplier', ''),
        '<<PRODUCTO_RECIBIDO>>': reception_data.get('product_name', ''),
        '<<CANTIDAD_RECIBIDA>>': str(reception_data.get('quantity', '')),
        '<<UNIDADES>>': reception_data.get('unit', ''),
        '<<PRODUCTO_REVISADO>>': reception_data.get('units_reviewed', ''),
        '<<NO_LOTE>>': reception_data.get('lot_number', ''),
        '<<CADUCIDAD>>': expiration_date,
        '<<FECHA_CONTROL>>': reception_date,
        '<<ENV>>': 'X' if reception_data['product_type'] == 'packaging' else '',
        '<<MP>>': 'X' if reception_data['product_type'] == 'raw_material' else '',
    }
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    
    # Replace in tables
    for table in doc.tables:
        replace_text_in_table(table, replacements)
    
    # Generate filename: YYYYMMDD LOT. <lote> <producto> <proveedor>
    date_str = datetime.strptime(reception_data['reception_date'], '%Y-%m-%d').strftime('%Y%m%d')
    lot_number = reception_data.get('lot_number', 'LOTE').replace('/', '_').replace('\\', '_')
    product_name = reception_data.get('product_name', 'Producto').replace('/', '_').replace('\\', '_')
    supplier = reception_data.get('supplier', 'Proveedor').replace('/', '_').replace('\\', '_')
    filename_base = f"{date_str} LOT. {lot_number} {product_name} {supplier}"
    
    docx_path = os.path.join(OUTPUT_DIR, f"{filename_base}.docx")
    pdf_path = os.path.join(OUTPUT_DIR, f"{filename_base}.pdf")
    
    # Save the filled document
    doc.save(docx_path)
    
    # Convert to PDF using LibreOffice headless mode (transparent, no UI)
    pdf_path = convert_docx_to_pdf_headless(docx_path, OUTPUT_DIR)
    
    return docx_path, pdf_path


def convert_docx_to_pdf_headless(docx_path, output_dir):
    """
    Convert DOCX to PDF using LibreOffice in headless mode (no UI).
    Falls back to sending DOCX if LibreOffice is not available.
    """
    import subprocess
    import shutil
    
    # Try LibreOffice paths on different platforms
    libreoffice_paths = [
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
        '/usr/bin/libreoffice',  # Linux
        '/usr/bin/soffice',  # Linux alternative
        'soffice',  # In PATH
    ]
    
    libreoffice = None
    for path in libreoffice_paths:
        if shutil.which(path) or os.path.exists(path):
            libreoffice = path
            break
    
    if not libreoffice:
        print("Warning: LibreOffice not found. PDF conversion skipped - will use DOCX.")
        return None
    
    try:
        # Run LibreOffice in headless mode to convert
        result = subprocess.run([
            libreoffice,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_path
        ], capture_output=True, text=True, timeout=60)
        
        # Check if PDF was created
        pdf_path = docx_path.replace('.docx', '.pdf')
        if os.path.exists(pdf_path):
            return pdf_path
        else:
            print(f"Warning: PDF not created. LibreOffice output: {result.stderr}")
            return None
            
    except subprocess.TimeoutExpired:
        print("Warning: LibreOffice conversion timed out.")
        return None
    except Exception as e:
        print(f"Warning: Could not convert to PDF: {e}")
        return None


def send_reception_email(pdf_path, reception_data, recipient=EMAIL_RECIPIENT):
    """
    Send the reception document via email using SMTP SSL
    """
    if not pdf_path or not os.path.exists(pdf_path):
        return False, "PDF file not found"
    
    try:
        # Create message
        msg = MIMEMultipart()
        msg['From'] = EMAIL_USER
        msg['To'] = recipient
        msg['Subject'] = f"Registro Recepción - {reception_data.get('product_name', 'Producto')} - Lote {reception_data.get('lot_number', '')}"
        
        # Email body
        body = f"""
Adjunto el registro de recepción:

- Producto: {reception_data.get('product_name', '')}
- Lote: {reception_data.get('lot_number', '')}
- Proveedor: {reception_data.get('supplier', '')}
- Cantidad: {reception_data.get('quantity', '')} {reception_data.get('unit', '')}
- Fecha Recepción: {reception_data.get('reception_date', '')}

Este documento ha sido generado automáticamente.
"""
        msg.attach(MIMEText(body, 'plain'))
        
        # Attach PDF
        with open(pdf_path, 'rb') as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
        
        encoders.encode_base64(part)
        filename = os.path.basename(pdf_path)
        part.add_header('Content-Disposition', f'attachment; filename= {filename}')
        msg.attach(part)
        
        # Send email using SSL
        with smtplib.SMTP_SSL(SMTP_SERVER, SMTP_PORT) as server:
            server.login(EMAIL_USER, EMAIL_PASS)
            server.sendmail(EMAIL_USER, recipient, msg.as_string())
        
        return True, f"Email enviado a {recipient}"
        
    except Exception as e:
        return False, f"Error enviando email: {str(e)}"


def process_reception_document(reception_data):
    """
    Main function to generate and send reception document
    
    Returns:
        dict with 'success', 'message', 'docx_path', 'pdf_path'
    """
    try:
        docx_path, pdf_path = generate_reception_document(reception_data)
        
        result = {
            'success': True,
            'message': 'Documento generado correctamente',
            'docx_path': docx_path,
            'pdf_path': pdf_path
        }
        
        # Try to send email if PDF was generated
        if pdf_path:
            email_sent, email_message = send_reception_email(pdf_path, reception_data)
            result['email_sent'] = email_sent
            result['email_message'] = email_message
        
        return result
        
    except Exception as e:
        return {
            'success': False,
            'message': f'Error generando documento: {str(e)}',
            'docx_path': None,
            'pdf_path': None
        }
